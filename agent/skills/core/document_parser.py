# document_parser.py
"""
Core skill: parse PDF, DOCX, XLSX, CSV, TXT, JSON, YAML, Markdown and more.
All paths resolve relative to /app/memory/ (or absolute within /app/).
"""

import os
import json
import csv
from pathlib import Path

NAME = "document_parser"
SHORT_DOC = "Parse and extract text, tables, and metadata from PDF, DOCX, XLSX, CSV, and other file formats."
DOC = (
    "Parse and extract content from PDF, DOCX, XLSX, CSV, TXT, MD, JSON, YAML files stored under /app/. "
    "Paths are relative to /app/memory/ unless absolute. "
    "Returns: read(path)→full document text ready for analysis; metadata(path)→author/date/page count as text; "
    "extract_tables(path)→tab-separated table rows as text; "
    "ocr(path, lang?)→extract text from an image file using Tesseract OCR; lang defaults to 'eng', supports 'eng+fra', 'ita', 'spa', etc.; "
    "convert_to_text(path, outpath)→extracts text from any file and saves to /app/memory/knowledge/; "
    "create_text(filename, content)→save a plain text or markdown document to /app/memory/knowledge/; "
    "create_pdf(filename, content)→generate a PDF from plain text, saved to /app/memory/knowledge/; "
    "create_pdf_from_markdown(filename, markdown_text)→generate a styled PDF from markdown, saved to /app/memory/knowledge/; "
    "list_supported()→list of supported file extensions. "
    "ALL document creation functions (create_text, create_pdf, create_pdf_from_markdown) save to /app/memory/knowledge/ "
    "and return a clickable markdown download link — do NOT pass that return value back as content. "
    "CRITICAL CALL FORMAT: arg1=filename, arg2=THE FULL TEXT CONTENT — NOT a path, NOT a title only. "
    "Example: <skill:document_parser.create_pdf_from_markdown>report.pdf, # My Title\n\nFull body text goes here.\n\n## Section\nMore text.</skill:document_parser.create_pdf_from_markdown>"
)

# ── Lazy optional imports ──────────────────────────────────────────────────────

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pypdf import PdfReader as _PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as _DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

try:
    from PIL import Image as _PILImage
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

try:
    import pytesseract as _pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# ── Path helpers ───────────────────────────────────────────────────────────────

_BASE      = Path("/app")
_MEMORY    = Path("/app/memory")
_KNOWLEDGE = Path("/app/memory/knowledge")
_KNOWLEDGE.mkdir(parents=True, exist_ok=True)

# Base URL used to build download links returned to the UI.
# Override with the AGENT_BASE_URL env var if the agent is not on localhost.
_AGENT_BASE_URL = os.getenv("AGENT_BASE_URL", "http://localhost:8001").rstrip("/")

def _resolve(path: str) -> Path:
    """Resolve path safely within /app/. Raises ValueError on traversal."""
    p = Path(path)
    if not p.is_absolute():
        candidate = (_MEMORY / p).resolve()
        if not candidate.exists():
            candidate = (_BASE / p).resolve()
    else:
        candidate = p.resolve()
    if not str(candidate).startswith(str(_BASE.resolve())):
        raise ValueError(f"Path escapes /app/: {path}")
    return candidate

def _ext(p: Path) -> str:
    """Return the lowercase extension of a Path without the leading dot."""
    return p.suffix.lower().lstrip(".")

# ── PDF ────────────────────────────────────────────────────────────────────────

def _read_pdf(p: Path) -> str:
    if HAS_PDFPLUMBER:
        pages = []
        with pdfplumber.open(p) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(f"--- Page {i} ---\n{t}")
        return "\n\n".join(pages) or "[PDF has no extractable text]"
    if HAS_PYPDF:
        reader = _PdfReader(str(p))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            t = page.extract_text() or ""
            if t.strip():
                pages.append(f"--- Page {i} ---\n{t}")
        return "\n\n".join(pages) or "[PDF has no extractable text]"
    return "[PDF support not installed. Run: pip install pdfplumber]"

def _pdf_tables(p: Path) -> str:
    if not HAS_PDFPLUMBER:
        return "[pdfplumber not installed. Run: pip install pdfplumber]"
    out = []
    with pdfplumber.open(p) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            for j, table in enumerate(page.extract_tables(), 1):
                out.append(f"--- Page {i}, Table {j} ---")
                for row in table:
                    out.append(" | ".join(str(c or "").strip() for c in row))
    return "\n".join(out) or "[No tables found in PDF]"

def _pdf_meta(p: Path) -> dict:
    m = {"pages": 0, "has_tables": False}
    if HAS_PDFPLUMBER:
        with pdfplumber.open(p) as pdf:
            m["pages"] = len(pdf.pages)
            for page in pdf.pages:
                if page.extract_tables():
                    m["has_tables"] = True
                    break
    elif HAS_PYPDF:
        m["pages"] = len(_PdfReader(str(p)).pages)
    return m

# ── DOCX ───────────────────────────────────────────────────────────────────────

def _read_docx(p: Path) -> str:
    if not HAS_DOCX:
        return "[DOCX support not installed. Run: pip install python-docx]"
    doc = _DocxDocument(str(p))
    parts = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name if para.style else ""
        prefix = "# " if "Heading 1" in style else "## " if "Heading 2" in style else ""
        parts.append(f"{prefix}{para.text}")
    for i, table in enumerate(doc.tables, 1):
        parts.append(f"\n[Table {i}]")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts) or "[DOCX appears empty]"

def _docx_meta(p: Path) -> dict:
    if not HAS_DOCX:
        return {}
    doc = _DocxDocument(str(p))
    word_count = sum(len(para.text.split()) for para in doc.paragraphs)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables":     len(doc.tables),
        "word_count": word_count,
    }

# ── Spreadsheets ───────────────────────────────────────────────────────────────

def _read_xlsx(p: Path) -> str:
    if not HAS_PANDAS:
        return "[pandas not available]"
    xl = pd.ExcelFile(str(p))
    parts = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        parts.append(f"--- Sheet: {sheet} ({len(df)} rows × {len(df.columns)} cols) ---")
        parts.append(df.to_string(index=False, max_rows=100))
    return "\n\n".join(parts)

def _read_csv(p: Path) -> str:
    if HAS_PANDAS:
        df = pd.read_csv(str(p))
        return f"({len(df)} rows × {len(df.columns)} cols)\n{df.to_string(index=False, max_rows=100)}"
    with open(p, newline="", encoding="utf-8", errors="replace") as f:
        rows = list(csv.reader(f))
    if not rows:
        return "[Empty CSV]"
    header = " | ".join(rows[0])
    sample = "\n".join(" | ".join(r) for r in rows[1:51])
    return f"{header}\n{'-'*40}\n{sample}\n({len(rows)-1} data rows total)"

# ── Images (OCR) ───────────────────────────────────────────────────────────────

def _read_image(p: Path) -> str:
    if not HAS_PILLOW:
        return "[Image support not installed. Run: pip install Pillow]"
    if not HAS_PYTESSERACT:
        return "[OCR not installed. Run: pip install pytesseract]"
    img = _PILImage.open(p)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    text = _pytesseract.image_to_string(img).strip()
    if not text:
        return (
            f"[Image '{p.name}': no readable text found via OCR. "
            f"Consider adding a companion '{p.stem}.txt' description file.]"
        )
    return f"[OCR extracted from {p.name}]\n{text}"

# ── Text / Markup ──────────────────────────────────────────────────────────────

def _read_txt(p: Path) -> str:
    with open(p, encoding="utf-8", errors="replace") as f:
        return f.read()

def _read_json(p: Path) -> str:
    with open(p, encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)

def _read_yaml(p: Path) -> str:
    if not HAS_YAML:
        return _read_txt(p)
    with open(p, encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)

# ── Format router ──────────────────────────────────────────────────────────────

_READERS = {
    "pdf":      _read_pdf,
    "docx":     _read_docx,
    "doc":      _read_docx,
    "xlsx":     _read_xlsx,
    "xls":      _read_xlsx,
    "csv":      _read_csv,
    "txt":      _read_txt,
    "md":       _read_txt,
    "markdown": _read_txt,
    "json":     _read_json,
    "jsonl":    _read_txt,
    "yaml":     _read_yaml,
    "yml":      _read_yaml,
    "log":      _read_txt,
    "xml":      _read_txt,
    "html":     _read_txt,
    "htm":      _read_txt,
    "py":       _read_txt,
    "js":       _read_txt,
    "ts":       _read_txt,
    "sh":       _read_txt,
    "png":      _read_image,
    "jpg":      _read_image,
    "jpeg":     _read_image,
    "webp":     _read_image,
    "gif":      _read_image,
    "bmp":      _read_image,
    "tiff":     _read_image,
    "tif":      _read_image,
}

# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC SKILL FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def read(*args) -> str:
    """
    Extract and return all text from a document.
    Usage: read(path)
    Supports: PDF, DOCX, XLSX, CSV, TXT, MD, JSON, YAML, XML, HTML, code files.
    Path is relative to /app/memory/ or absolute within /app/.
    """
    try:
        if not args:
            return "Error: provide a file path.  Usage: read(path)"
        p = _resolve(str(args[0]).strip())
        if not p.exists():
            return f"File not found: {p}"
        reader = _READERS.get(_ext(p))
        if reader is None:
            try:
                return _read_txt(p)
            except Exception:
                return (
                    f"Unsupported file type: .{_ext(p)}. "
                    f"Supported: {', '.join(sorted(_READERS))}"
                )
        return reader(p)
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error reading '{args[0] if args else '?'}': {e}"


def extract_tables(*args) -> str:
    """
    Extract all tables from a document as formatted text.
    Usage: extract_tables(path)
    Best results with: PDF (requires pdfplumber), DOCX, XLSX, CSV.
    """
    try:
        if not args:
            return "Error: provide a file path"
        p = _resolve(str(args[0]).strip())
        if not p.exists():
            return f"File not found: {p}"
        ext = _ext(p)
        if ext == "pdf":
            return _pdf_tables(p)
        if ext in ("docx", "doc"):
            if not HAS_DOCX:
                return "[python-docx not installed. Run: pip install python-docx]"
            doc = _DocxDocument(str(p))
            if not doc.tables:
                return "[No tables found in document]"
            out = []
            for i, table in enumerate(doc.tables, 1):
                out.append(f"--- Table {i} ---")
                for row in table.rows:
                    out.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(out)
        if ext in ("xlsx", "xls"):
            return _read_xlsx(p)
        if ext == "csv":
            return _read_csv(p)
        return f"Table extraction not supported for .{ext} files"
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error: {e}"


def metadata(*args) -> str:
    """
    Return file metadata: type, size, page count, word count, columns, etc.
    Usage: metadata(path)
    """
    try:
        if not args:
            return "Error: provide a file path"
        p = _resolve(str(args[0]).strip())
        if not p.exists():
            return f"File not found: {p}"
        stat = p.stat()
        ext  = _ext(p)
        info = {
            "file":       p.name,
            "path":       str(p),
            "extension":  ext,
            "size_bytes": stat.st_size,
            "size_kb":    round(stat.st_size / 1024, 2),
        }
        if ext == "pdf":
            info.update(_pdf_meta(p))
        elif ext in ("docx", "doc") and HAS_DOCX:
            info.update(_docx_meta(p))
        elif ext in ("xlsx", "xls") and HAS_PANDAS:
            xl = pd.ExcelFile(str(p))
            info["sheets"]      = xl.sheet_names
            info["sheet_count"] = len(xl.sheet_names)
        elif ext == "csv" and HAS_PANDAS:
            df = pd.read_csv(str(p), nrows=0)
            full = pd.read_csv(str(p))
            info["rows"]    = len(full)
            info["columns"] = list(df.columns)
        elif ext in _READERS:
            try:
                text = _read_txt(p)
                info["lines"]      = text.count("\n") + 1
                info["word_count"] = len(text.split())
                info["char_count"] = len(text)
            except Exception:
                pass
        return json.dumps(info, indent=2)
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error: {e}"


def summarize(*args) -> str:
    """
    Extract text from a file and truncate to a safe size for LLM context.
    Usage: summarize(path) or summarize(path, max_chars)
    Default max_chars: 4000.
    """
    try:
        if not args:
            return "Error: provide a file path"
        p        = _resolve(str(args[0]).strip())
        max_chars = int(args[1]) if len(args) > 1 else 4000
        if not p.exists():
            return f"File not found: {p}"
        reader   = _READERS.get(_ext(p), _read_txt)
        full     = reader(p)
        total    = len(full)
        if total <= max_chars:
            return full
        return (
            f"{full[:max_chars]}\n\n"
            f"[... truncated — {total} total chars, showing first {max_chars}. "
            f"Use read() for full content.]"
        )
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error: {e}"


def convert_to_text(*args) -> str:
    """
    Extract text from any supported document and save it as a .txt file in /app/memory/knowledge/.
    Usage: convert_to_text(source_path) or convert_to_text(source_path, output_filename)
    Output always written to /app/memory/knowledge/.
    """
    try:
        if not args:
            return "Error: provide a source file path"
        p = _resolve(str(args[0]).strip())
        if not p.exists():
            return f"File not found: {p}"
        reader   = _READERS.get(_ext(p), _read_txt)
        text     = reader(p)
        out_name = str(args[1]).strip() if len(args) > 1 else f"{p.stem}_extracted.txt"
        out_path = _KNOWLEDGE / out_name
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        url = f"{_AGENT_BASE_URL}/download/{out_path.name}"
        return f"Saved {len(text)} chars — [{out_path.name}]({url})"
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error: {e}"


def create_text(filename, content) -> str:
    """
    Write a plain-text or markdown document and save it to /app/memory/knowledge/.

    ARGUMENTS:
      filename — e.g. "notes.txt" or "report.md"
      content  — THE COMPLETE TEXT CONTENT of the document.

    CORRECT usage:
      <skill:document_parser.create_text>notes.txt, Line one of content.

    Line two here.</skill:document_parser.create_text>

    Returns a markdown download link. Do NOT pass it back as content.
    """
    try:
        if not filename or not content:
            return "Error: provide filename and content.  Usage: create_text(filename, content)"
        filename = str(filename).strip()
        content  = str(content)
        out_path = _KNOWLEDGE / filename
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(content, encoding="utf-8")
        url = f"{_AGENT_BASE_URL}/download/{filename}"
        return f"File created: [{filename}]({url})"
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error creating file: {e}"


def create_pdf(filename, content) -> str:
    """
    Generate a PDF from plain text and save it to /app/memory/.

    ARGUMENTS:
      filename — e.g. "report.pdf"
      content  — THE COMPLETE PLAIN TEXT — all paragraphs, not just a title.
                 Blank lines = paragraph breaks.

    CORRECT usage:
      <skill:document_parser.create_pdf>report.pdf, First paragraph text here.

    Second paragraph here.</skill:document_parser.create_pdf>

    Returns a markdown download link: [filename.pdf](download URL)
    Do NOT pass this return value back as content.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return "[PDF generation not available. Run: pip install reportlab]"
    try:
        if not filename or not content:
            return "Error: provide filename and content.  Usage: create_pdf(filename, content)"
        filename = str(filename).strip()
        content  = str(content)
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
        out_path = _KNOWLEDGE / filename
        out_path.parent.mkdir(parents=True, exist_ok=True)

        styles  = getSampleStyleSheet()
        body    = styles["BodyText"]
        doc     = SimpleDocTemplate(
            str(out_path),
            pagesize=LETTER,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch,  bottomMargin=inch,
        )
        story = []
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para.replace("\n", "<br/>"), body))
                story.append(Spacer(1, 0.15 * inch))
        if not story:
            story.append(Paragraph("(empty)", body))
        doc.build(story)
        url = f"{_AGENT_BASE_URL}/download/{filename}"
        return f"PDF created: [{filename}]({url})"
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error creating PDF: {e}"


def create_pdf_from_markdown(filename, markdown_text) -> str:
    """
    Generate a styled PDF from markdown-like text and save it to /app/memory/.

    ARGUMENTS:
      filename      — e.g. "report.pdf"
      markdown_text — THE COMPLETE DOCUMENT TEXT including all headings,
                      paragraphs, bullets, etc. NOT just a title. NOT a file path.

    Supported markdown: # Heading1, ## Heading2, ### Heading3,
                        - bullet / * bullet, **bold**, blank line = paragraph break.

    CORRECT skill tag example:
      <skill:document_parser.create_pdf_from_markdown>report.pdf, # Report Title

    ## Introduction
    This section explains the purpose of the report.

    ## Details
    - Point one
    - Point two
    </skill:document_parser.create_pdf_from_markdown>

    Returns a markdown download link: [filename.pdf](download URL)
    Do NOT pass this return value as the content argument in another call.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return "[PDF generation not available. Run: pip install reportlab]"
    try:
        if not filename or not markdown_text:
            return "Error: provide filename and markdown_text.  Usage: create_pdf_from_markdown(filename, text)"
        filename = str(filename).strip()
        text     = str(markdown_text)
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
        out_path = _KNOWLEDGE / filename
        out_path.parent.mkdir(parents=True, exist_ok=True)

        import re
        base_styles = getSampleStyleSheet()
        h1 = ParagraphStyle("H1", parent=base_styles["Heading1"], fontSize=18, spaceAfter=8)
        h2 = ParagraphStyle("H2", parent=base_styles["Heading2"], fontSize=14, spaceAfter=6)
        h3 = ParagraphStyle("H3", parent=base_styles["Heading3"], fontSize=12, spaceAfter=4)
        body   = base_styles["BodyText"]
        bullet = ParagraphStyle("Bullet", parent=body, leftIndent=20, bulletIndent=10)

        def _inline(line: str) -> str:
            """Convert **bold** → <b>bold</b> for reportlab."""
            return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)

        doc   = SimpleDocTemplate(
            str(out_path),
            pagesize=LETTER,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch,  bottomMargin=inch,
        )
        story     = []
        buf_lines = []

        def _flush_buf():
            if buf_lines:
                para = " ".join(buf_lines).strip()
                if para:
                    story.append(Paragraph(_inline(para), body))
                    story.append(Spacer(1, 0.1 * inch))
                buf_lines.clear()

        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if line.startswith("### "):
                _flush_buf()
                story.append(Paragraph(_inline(line[4:]), h3))
                story.append(Spacer(1, 0.05 * inch))
            elif line.startswith("## "):
                _flush_buf()
                story.append(Paragraph(_inline(line[3:]), h2))
                story.append(Spacer(1, 0.08 * inch))
            elif line.startswith("# "):
                _flush_buf()
                story.append(Paragraph(_inline(line[2:]), h1))
                story.append(Spacer(1, 0.1 * inch))
            elif re.match(r"^[-*] ", line):
                _flush_buf()
                story.append(Paragraph(_inline(line[2:]), bullet, bulletText="•"))
                story.append(Spacer(1, 0.05 * inch))
            elif line == "":
                _flush_buf()
            else:
                buf_lines.append(line)
        _flush_buf()

        if not story:
            story.append(Paragraph("(empty)", body))
        doc.build(story)
        url = f"{_AGENT_BASE_URL}/download/{filename}"
        return f"PDF created: [{filename}]({url})"
    except ValueError as e:
        return f"Security error: {e}"
    except Exception as e:
        return f"Error creating PDF: {e}"


def ocr(path: str, lang: str = "eng") -> str:
    """
    Extract text from an image file using Tesseract OCR.
    Works on screenshots, scanned documents, receipts, and any image with printed text.

    Args:
        path: Path to the image (PNG, JPG, TIFF, BMP, WebP, etc.)
        lang: Tesseract language code. Default 'eng'. Combine with '+': 'eng+fra', 'eng+srp_latn'.
              Common codes: eng, ita, spa, fra, deu, ell, srp, srp_latn.

    Returns:
        Extracted text as a plain string, or an error message prefixed with ❌.
    """
    if not HAS_PILLOW:
        return "❌ OCR unavailable — Pillow not installed. Run: pip install Pillow"
    if not HAS_PYTESSERACT:
        return "❌ OCR unavailable — pytesseract not installed. Run: pip install pytesseract (and apt-get install tesseract-ocr)"
    try:
        p = _resolve(path)
        if not p.exists():
            return f"❌ File not found: {path}"
        img = _PILImage.open(p)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = _pytesseract.image_to_string(img, lang=lang).strip()
        return text if text else "(no text detected in image)"
    except _pytesseract.TesseractNotFoundError:
        return "❌ Tesseract binary not found. Install: apt-get install -y tesseract-ocr"
    except ValueError as e:
        return f"❌ Security error: {e}"
    except Exception as e:
        return f"❌ OCR error: {e}"


def list_supported(*args) -> str:
    """
    List all supported file formats and whether optional libraries are installed.
    Usage: list_supported()
    """
    status = {
        "PDF":           ("✅ pdfplumber"         if HAS_PDFPLUMBER else
                          "⚠️ pypdf (fallback)"    if HAS_PYPDF      else
                          "❌ pip install pdfplumber"),
        "DOCX / DOC":    "✅ python-docx"          if HAS_DOCX       else "❌ pip install python-docx",
        "XLSX / XLS":    "✅ pandas"                if HAS_PANDAS     else "❌ pip install pandas",
        "CSV":           "✅ built-in (+ pandas)",
        "TXT / MD / LOG":"✅ built-in",
        "JSON / JSONL":  "✅ built-in",
        "YAML / YML":    "✅ pyyaml"                if HAS_YAML       else "⚠️ fallback to raw text",
        "XML / HTML":    "✅ read as plain text",
        "PY / JS / TS":  "✅ read as plain text",
        "PNG/JPG/WEBP…": ("✅ pytesseract OCR" if HAS_PYTESSERACT and HAS_PILLOW else
                          "❌ pip install Pillow pytesseract (+ tesseract-ocr system package)"),
    }
    lines = ["document_parser — supported formats\n"]
    for fmt, s in status.items():
        lines.append(f"  {fmt:<20} {s}")
    status["PDF generation"] = "✅ reportlab" if HAS_REPORTLAB else "❌ pip install reportlab"
    lines.append(
        "\nFunctions: read, extract_tables, metadata, summarize, "
        "convert_to_text, create_pdf, create_pdf_from_markdown, list_supported"
    )
    return "\n".join(lines)
